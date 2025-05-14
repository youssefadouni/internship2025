import os
import logging
import pytesseract
import pdfplumber
from docx import Document
from openpyxl import load_workbook
from PIL import Image
import concurrent.futures
import time
import tqdm
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from collections import Counter
from langdetect import detect
from textblob import TextBlob
import argparse
import sys
import csv
import json

# Set up logging with more detailed configuration
log_formatter = logging.Formatter('%(asctime)s - %(levelname)s - [%(filename)s:%(lineno)d] - %(message)s')
log_file = logging.FileHandler("extraction.log")
log_file.setFormatter(log_formatter)
log_console = logging.StreamHandler()
log_console.setFormatter(log_formatter)

logger = logging.getLogger('text_extractor')
logger.setLevel(logging.INFO)
logger.addHandler(log_file)
logger.addHandler(log_console)

# Constants
UPLOADS_DIR = "uploads"
OUTPUT_DIR = "extracted_texts"
MAX_WORKERS = 4  # Number of parallel workers for multi-threading

# Ensure NLTK resources are available
def download_nltk_resources():
    """Download required NLTK resources if not already present"""
    required_resources = [
        'tokenizers/punkt',
        'corpora/stopwords',
        'tokenizers/punkt/punkt.pickle',
        'tokenizers/punkt/PY3/english.pickle'
    ]
    
    for resource in required_resources:
        try:
            nltk.data.find(resource)
            logger.debug(f"NLTK resource '{resource}' already available")
        except LookupError:
            logger.info(f"Downloading NLTK resource: {resource}")
            # Extract the main package name from the resource path
            if 'punkt' in resource:
                nltk.download('punkt')
            elif 'stopwords' in resource:
                nltk.download('stopwords')
            else:
                # For any other resources we might add in the future
                package = resource.split('/')[0]
                nltk.download(package)
    
    logger.info("All required NLTK resources are available")

class TextExtractor:
    """Class to handle text extraction from various file formats"""
    
    def __init__(self, input_dir=UPLOADS_DIR, output_dir=OUTPUT_DIR, max_workers=MAX_WORKERS):
        self.input_dir = input_dir
        self.output_dir = output_dir
        self.max_workers = max_workers
        self.supported_formats = {
            ".pdf": self.extract_text_from_pdf,
            ".docx": self.extract_text_from_docx,
            ".xlsx": self.extract_text_from_xlsx,
            ".txt": self.extract_text_from_txt,
            ".csv": self.extract_text_from_csv,
            ".json": self.extract_text_from_json
        }
        
        # Create output directory if it doesn't exist
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
            logger.info(f"Created output directory: {self.output_dir}")
    
    def extract_text_from_pdf(self, file_path):
        """Extract text from PDF files with OCR fallback for scanned documents.
        
        First attempts direct text extraction, then falls back to OCR if no text found.
        OCR is more resource-intensive but handles scanned documents.
        """
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                # First pass: Try native text extraction (faster, works for digital PDFs)
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                
                # Second pass: If no text found, PDF might be scanned/image-based
                if not text.strip():
                    logger.info(f"No text directly extracted from {file_path}. Attempting OCR.")
                    full_ocr_text = ""
                    
                    # Show progress for OCR which can be time-consuming
                    for i, page in enumerate(tqdm.tqdm(pdf.pages, desc="OCR Processing")):
                        try:
                            # Convert page to image at 300dpi for OCR
                            im = page.to_image(resolution=300)
                            page_ocr_text = pytesseract.image_to_string(im.original, lang='eng')
                            if page_ocr_text:
                                full_ocr_text += page_ocr_text + "\n"
                        except Exception as ocr_page_error:
                            logger.error(f"Error during OCR for page {i+1} of {file_path}: {ocr_page_error}")
                    text = full_ocr_text
        except Exception as e:
            logger.error(f"Error processing PDF file {file_path}: {e}", exc_info=True)
        return text.strip()

    def extract_text_from_docx(self, file_path):
        """Extract text from Word documents (.docx).
        
        Extracts text from paragraphs, tables, headers, and footers.
        """
        text = ""
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            # Extract text from headers and footers
            for section in doc.sections:
                for header in section.header.paragraphs:
                    text += header.text + "\n"
                for footer in section.footer.paragraphs:
                    text += footer.text + "\n"
                    
        except Exception as e:
            logger.error(f"Error processing DOCX file {file_path}: {e}", exc_info=True)
        return text.strip()

    def extract_text_from_xlsx(self, file_path):
        """Extract text from Excel spreadsheets (.xlsx).
        
        Uses read_only and data_only modes for better performance with large files.
        Preserves sheet and row structure.
        """
        text = ""
        try:
            workbook = load_workbook(filename=file_path, read_only=True, data_only=True)
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text += f"\n--- Sheet: {sheet_name} ---\n"
                for row in sheet.iter_rows():
                    row_text = ""
                    for cell in row:
                        if cell.value is not None:
                            row_text += str(cell.value) + "\t"
                    if row_text:
                        text += row_text + "\n"
        except Exception as e:
            logger.error(f"Error processing XLSX file {file_path}: {e}", exc_info=True)
        return text.strip()
    
    def extract_text_from_txt(self, file_path):
        """Extract text from plain text files (.txt)."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except UnicodeDecodeError:
            # Try different encodings if UTF-8 fails
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read().strip()
            except Exception as e:
                logger.error(f"Error processing TXT file {file_path}: {e}", exc_info=True)
                return ""
        except Exception as e:
            logger.error(f"Error processing TXT file {file_path}: {e}", exc_info=True)
            return ""
    
    def extract_text_from_csv(self, file_path):
        """Extract text from CSV files (.csv)."""
        text = ""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                csv_reader = csv.reader(file)
                for row in csv_reader:
                    text += "\t".join(row) + "\n"
        except UnicodeDecodeError:
            # Try different encodings if UTF-8 fails
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    csv_reader = csv.reader(file)
                    for row in csv_reader:
                        text += "\t".join(row) + "\n"
            except Exception as e:
                logger.error(f"Error processing CSV file {file_path}: {e}", exc_info=True)
        except Exception as e:
            logger.error(f"Error processing CSV file {file_path}: {e}", exc_info=True)
        return text.strip()
    
    def extract_text_from_json(self, file_path):
        """Extract text from JSON files (.json)."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
                # Convert JSON to formatted string
                return json.dumps(data, indent=2)
        except Exception as e:
            logger.error(f"Error processing JSON file {file_path}: {e}", exc_info=True)
            return ""
    
    def process_file(self, file_path):
        """Process a single file and return its extracted text."""
        filename = os.path.basename(file_path)
        file_ext = os.path.splitext(filename)[1].lower()
        
        if file_ext in self.supported_formats:
            logger.info(f"Processing file: {filename}")
            start_time = time.time()
            
            try:
                extracted_text = self.supported_formats[file_ext](file_path)
                processing_time = time.time() - start_time
                
                if extracted_text:
                    logger.info(f"Successfully extracted text from {filename} in {processing_time:.2f} seconds")
                    # Save individual file text to output directory
                    output_path = os.path.join(self.output_dir, f"{os.path.splitext(filename)[0]}.txt")
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(extracted_text)
                    
                    return {
                        "filename": filename,
                        "text": extracted_text,
                        "processing_time": processing_time
                    }
                else:
                    logger.warning(f"No text extracted from {filename}")
                    return {"filename": filename, "text": "", "error": "No text extracted"}
            except Exception as e:
                logger.error(f"Error processing {filename}: {e}", exc_info=True)
                return {"filename": filename, "text": "", "error": str(e)}
        else:
            logger.info(f"Skipping unsupported file type: {filename}")
            return {"filename": filename, "text": "", "error": "Unsupported file type"}
    
    def process_files(self):
        """Process all supported files in the input directory using multi-threading."""
        if not os.path.exists(self.input_dir):
            logger.error(f"Directory not found: {self.input_dir}")
            return []
        
        if not os.listdir(self.input_dir):
            logger.info(f"No files found in {self.input_dir} directory.")
            return []
        
        file_paths = []
        for filename in os.listdir(self.input_dir):
            file_path = os.path.join(self.input_dir, filename)
            if os.path.isfile(file_path):
                file_paths.append(file_path)
        
        results = []
        if file_paths:
            logger.info(f"Starting multi-threaded processing of {len(file_paths)} files with {self.max_workers} workers")
            
            # Process files in parallel using ThreadPoolExecutor
            with concurrent.futures.ThreadPoolExecutor(max_workers=self.max_workers) as executor:
                # Submit all tasks and create a dictionary mapping futures to filenames for progress tracking
                future_to_file = {executor.submit(self.process_file, file_path): file_path 
                                for file_path in file_paths}
                
                # Process results as they complete with progress bar
                for future in tqdm.tqdm(concurrent.futures.as_completed(future_to_file), 
                                        total=len(file_paths), 
                                        desc="Extracting Text"):
                    file_path = future_to_file[future]
                    try:
                        result = future.result()
                        results.append(result)
                    except Exception as e:
                        filename = os.path.basename(file_path)
                        logger.error(f"Unhandled exception processing {filename}: {e}", exc_info=True)
                        results.append({"filename": filename, "text": "", "error": str(e)})
        
        return results

class TextAnalyzer:
    """Class to perform text analysis on extracted content"""
    
    def __init__(self):
        # Ensure we have the required NLTK resources
        try:
            download_nltk_resources()
            # Initialize stopwords - with fallback if resource is missing
            try:
                self.stop_words = set(stopwords.words('english'))
            except Exception as e:
                logger.warning(f"Could not load stopwords: {e}. Using a minimal set instead.")
                # Fallback to a minimal set of common English stopwords
                self.stop_words = set(['i', 'me', 'my', 'myself', 'we', 'our', 'ours', 'ourselves', 'you', 'your', 
                                      'yours', 'yourself', 'yourselves', 'he', 'him', 'his', 'himself', 'she', 
                                      'her', 'hers', 'herself', 'it', 'its', 'itself', 'they', 'them', 'their', 
                                      'theirs', 'themselves', 'what', 'which', 'who', 'whom', 'this', 'that', 
                                      'these', 'those', 'am', 'is', 'are', 'was', 'were', 'be', 'been', 'being', 
                                      'have', 'has', 'had', 'having', 'do', 'does', 'did', 'doing', 'a', 'an', 
                                      'the', 'and', 'but', 'if', 'or', 'because', 'as', 'until', 'while', 'of', 
                                      'at', 'by', 'for', 'with', 'about', 'against', 'between', 'into', 'through', 
                                      'during', 'before', 'after', 'above', 'below', 'to', 'from', 'up', 'down', 
                                      'in', 'out', 'on', 'off', 'over', 'under', 'again', 'further', 'then', 
                                      'once', 'here', 'there', 'when', 'where', 'why', 'how', 'all', 'any', 
                                      'both', 'each', 'few', 'more', 'most', 'other', 'some', 'such', 'no', 
                                      'nor', 'not', 'only', 'own', 'same', 'so', 'than', 'too', 'very'])
        except Exception as e:
            logger.error(f"Error initializing TextAnalyzer: {e}", exc_info=True)
            # Initialize with empty stopwords set as fallback
            self.stop_words = set()
    
    def extract_keywords(self, text, top_n=10):
        """Extract the most frequent meaningful words from text."""
        if not text:
            return []
        
        try:    
            # Tokenize and convert to lowercase
            words = word_tokenize(text.lower())
            
            # Remove stopwords, punctuation, and short words
            filtered_words = [word for word in words if word.isalpha() and 
                             word not in self.stop_words and len(word) > 2]
            
            # Count word frequencies
            word_freq = Counter(filtered_words)
            
            # Return top N keywords
            return word_freq.most_common(top_n)
        except Exception as e:
            logger.error(f"Error extracting keywords: {e}", exc_info=True)
            return []
    
    def detect_language(self, text):
        """Detect the language of the text."""
        if not text or len(text.strip()) < 20:  # Need sufficient text for reliable detection
            return "Unknown (insufficient text)"
            
        try:
            lang_code = detect(text)
            # Map language codes to full names (could be expanded)
            language_map = {
                'en': 'English',
                'fr': 'French',
                'es': 'Spanish',
                'de': 'German',
                'it': 'Italian',
                'pt': 'Portuguese',
                'nl': 'Dutch',
                'ru': 'Russian',
                'zh-cn': 'Chinese (Simplified)',
                'zh-tw': 'Chinese (Traditional)',
                'ja': 'Japanese',
                'ko': 'Korean',
                'ar': 'Arabic'
            }
            return language_map.get(lang_code, f"Other ({lang_code})")
        except Exception as e:
            logger.error(f"Language detection error: {e}", exc_info=True)
            return "Unknown (detection error)"
    
    def analyze_sentiment(self, text):
        """Analyze the sentiment of the text (positive, negative, neutral)."""
        if not text:
            return {"polarity": 0, "sentiment": "Neutral", "confidence": 0}
            
        try:
            analysis = TextBlob(text)
            polarity = analysis.sentiment.polarity
            
            # Determine sentiment category and confidence
            if polarity > 0.1:
                sentiment = "Positive"
                confidence = min(abs(polarity) * 10, 1.0)  # Scale to 0-1
            elif polarity < -0.1:
                sentiment = "Negative"
                confidence = min(abs(polarity) * 10, 1.0)  # Scale to 0-1
            else:
                sentiment = "Neutral"
                confidence = 1.0 - (abs(polarity) * 10)  # Higher confidence for values closer to 0
                
            return {
                "polarity": round(polarity, 2),
                "sentiment": sentiment,
                "confidence": round(confidence, 2)
            }
        except Exception as e:
            logger.error(f"Sentiment analysis error: {e}", exc_info=True)
            return {"polarity": 0, "sentiment": "Unknown", "confidence": 0}
    
    def get_text_stats(self, text):
        """Get basic statistics about the text."""
        if not text:
            return {
                "word_count": 0,
                "character_count": 0,
                "sentence_count": 0,
                "avg_word_length": 0,
                "avg_sentence_length": 0
            }
        
        try:    
            # Count words, characters, and sentences
            words = word_tokenize(text)
            word_count = len(words)
            char_count = len(text)
            sentences = text.split('.')
            sentence_count = len([s for s in sentences if s.strip()])
            
            # Calculate averages
            avg_word_length = char_count / word_count if word_count > 0 else 0
            avg_sentence_length = word_count / sentence_count if sentence_count > 0 else 0
            
            return {
                "word_count": word_count,
                "character_count": char_count,
                "sentence_count": sentence_count,
                "avg_word_length": round(avg_word_length, 2),
                "avg_sentence_length": round(avg_sentence_length, 2)
            }
        except Exception as e:
            logger.error(f"Error calculating text statistics: {e}", exc_info=True)
            return {
                "word_count": 0,
                "character_count": len(text),
                "sentence_count": 0,
                "avg_word_length": 0,
                "avg_sentence_length": 0,
                "error": str(e)
            }
    
    def analyze_text(self, text):
        """Perform comprehensive analysis on the text."""
        result = {}
        
        # Try each analysis component separately to ensure one failure doesn't stop everything
        try:
            result["keywords"] = self.extract_keywords(text)
        except Exception as e:
            logger.error(f"Keyword extraction failed: {e}", exc_info=True)
            result["keywords"] = []
            
        try:
            result["language"] = self.detect_language(text)
        except Exception as e:
            logger.error(f"Language detection failed: {e}", exc_info=True)
            result["language"] = "Unknown (error during detection)"
            
        try:
            result["sentiment"] = self.analyze_sentiment(text)
        except Exception as e:
            logger.error(f"Sentiment analysis failed: {e}", exc_info=True)
            result["sentiment"] = {"polarity": 0, "sentiment": "Unknown", "confidence": 0}
            
        try:
            result["stats"] = self.get_text_stats(text)
        except Exception as e:
            logger.error(f"Text statistics calculation failed: {e}", exc_info=True)
            result["stats"] = {"word_count": 0, "character_count": len(text), "sentence_count": 0, 
                           "avg_word_length": 0, "avg_sentence_length": 0}
        
        return result

def main():
    # Set up command line argument parsing
    parser = argparse.ArgumentParser(description="Enhanced Text Extraction Tool")
    parser.add_argument("-i", "--input", default=UPLOADS_DIR, help=f"Input directory (default: {UPLOADS_DIR})")
    parser.add_argument("-o", "--output", default=OUTPUT_DIR, help=f"Output directory (default: {OUTPUT_DIR})")
    parser.add_argument("-w", "--workers", type=int, default=MAX_WORKERS, help=f"Number of worker threads (default: {MAX_WORKERS})")
    parser.add_argument("-v", "--verbose", action="store_true", help="Enable verbose logging")
    parser.add_argument("-a", "--analyze", action="store_true", help="Perform text analysis on extracted content")
    args = parser.parse_args()
    
    # Download required NLTK resources at startup
    logger.info("Checking for required NLTK resources...")
    try:
        download_nltk_resources()
    except Exception as e:
        logger.error(f"Error downloading NLTK resources: {e}. Some text analysis features may not work properly.")
        logger.info("Continuing with extraction process despite NLTK resource issues.")

    
    # Set logging level based on verbosity flag
    if args.verbose:
        logger.setLevel(logging.DEBUG)
        logger.debug("Verbose logging enabled")
    
    # Check for Tesseract installation
    try:
        tesseract_version = pytesseract.get_tesseract_version()
        logger.info(f"Tesseract version {tesseract_version} found.")
    except pytesseract.TesseractNotFoundError:
        logger.error("Tesseract is not installed or not found in your PATH.")
        logger.error("Please install Tesseract and ensure it's in your system's PATH.")
        sys.exit(1)
    
    # Start the extraction process
    logger.info("Starting enhanced text extraction process...")
    start_time = time.time()
    
    # Initialize and run the text extractor
    extractor = TextExtractor(input_dir=args.input, output_dir=args.output, max_workers=args.workers)
    results = extractor.process_files()
    
    # Perform text analysis if requested
    if args.analyze and results:
        logger.info("Performing text analysis on extracted content...")
        analyzer = TextAnalyzer()
        
        # Create analysis output directory
        analysis_dir = os.path.join(args.output, "analysis")
        if not os.path.exists(analysis_dir):
            os.makedirs(analysis_dir)
        
        # Analyze each document and save results
        for result in tqdm.tqdm(results, desc="Analyzing Text"):
            if result.get("text"):
                analysis = analyzer.analyze_text(result["text"])
                
                # Save analysis results as JSON
                analysis_file = os.path.join(analysis_dir, f"{os.path.splitext(result['filename'])[0]}_analysis.json")
                with open(analysis_file, 'w', encoding='utf-8') as f:
                    json.dump({
                        "filename": result["filename"],
                        "analysis": analysis
                    }, f, indent=2)
    
    # Calculate and display summary
    total_time = time.time() - start_time
    successful = sum(1 for r in results if r.get("text"))
    failed = sum(1 for r in results if r.get("error"))
    
    print("\n" + "=" * 50)
    print(f"Text Extraction Summary")
    print("=" * 50)
    print(f"Total files processed: {len(results)}")
    print(f"Successfully extracted: {successful}")
    print(f"Failed to extract: {failed}")
    print(f"Total processing time: {total_time:.2f} seconds")
    print("=" * 50)
    
    # Display file-specific results
    if results:
        print("\nFile Details:")
        for result in results:
            status = "✓" if result.get("text") else "✗"
            filename = result["filename"]
            if result.get("text"):
                text_preview = result["text"][:100] + "..." if len(result["text"]) > 100 else result["text"]
                print(f"{status} {filename} - {len(result['text'])} chars")
            else:
                error = result.get("error", "Unknown error")
                print(f"{status} {filename} - Error: {error}")
    
    logger.info(f"Enhanced text extraction process completed in {total_time:.2f} seconds")

if __name__ == "__main__":
    main()