import os
import pytesseract
from PIL import Image
import re
from pypdf import PdfReader
from io import BytesIO
from docx import Document
from openpyxl import load_workbook

def clean_extracted_text(text):
    """Clean OCR-extracted text by fixing common issues"""
    text = re.sub(r'(\w)\1{2,}', r'\1', text)
    text = re.sub(r'\s+', ' ', text)
    text = text.replace('|', 'I').replace('[', '').replace(']', '')
    return text.strip()

def extract_pdf_content(pdf_path, output_dir):
    base_name = os.path.basename(pdf_path)
    txt_filename = os.path.splitext(base_name)[0] + ".txt"
    txt_path = os.path.join(output_dir, txt_filename)

    with open(txt_path, 'w', encoding='utf-8') as txt_file:
        reader = PdfReader(pdf_path)

        for page_num, page in enumerate(reader.pages, start=1):
            txt_file.write(f"\n=== Page {page_num} ===\n")
            
            text_content = page.extract_text()
            
            if hasattr(page, 'images') and page.images:
                if text_content:
                    txt_file.write(text_content)
                
                for img_num, image in enumerate(page.images, start=1):
                    try:
                        img = Image.open(BytesIO(image.data))
                        if img.mode != 'L':
                            img = img.convert('L')
                        
                        custom_config = r'--oem 3 --psm 6'
                        ocr_text = pytesseract.image_to_string(img, config=custom_config)
                        cleaned_text = clean_extracted_text(ocr_text)
                        
                        if cleaned_text.strip():
                            txt_file.write(f"\n[IMAGE TEXT - {img_num}]: {cleaned_text}\n")
                    except Exception as e:
                        txt_file.write(f"\n[IMAGE PROCESSING ERROR: {str(e)}]\n")
            else:
                if text_content:
                    txt_file.write(text_content)
                else:
                    txt_file.write("\n[No extractable content]\n")

    return txt_path

def extract_docx_content(docx_path, output_dir):
    """Extracts text and tables from DOCX while preserving original order"""
    from docx import Document
    
    base_name = os.path.basename(docx_path)
    txt_filename = os.path.splitext(base_name)[0] + ".txt"
    txt_path = os.path.join(output_dir, txt_filename)

    doc = Document(docx_path)
    
    with open(txt_path, 'w', encoding='utf-8') as txt_file:
        current_table_index = 0
        
        for element in doc.element.body:
            if element.tag.endswith('p'):
                para_text = ''
                for run in element.iterchildren():
                    if run.text:
                        para_text += run.text
                if para_text.strip():
                    txt_file.write(para_text + '\n')
            
            elif element.tag.endswith('tbl'):
                try:
                    if current_table_index < len(doc.tables):
                        table = doc.tables[current_table_index]
                        current_table_index += 1
                        
                        txt_file.write('\n[TABLE START]\n')
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                cell_text = ' '.join(
                                    p.text for p in cell.paragraphs 
                                    if p.text and p.text.strip()
                                )
                                row_text.append(cell_text)
                            txt_file.write('\t'.join(row_text) + '\n')
                        txt_file.write('[TABLE END]\n\n')
                    else:
                        txt_file.write('\n[TABLE FORMAT ERROR]\n')
                except Exception as e:
                    txt_file.write(f'\n[TABLE PROCESSING ERROR: {str(e)}]\n')
                    continue

    return txt_path

def extract_xlsx_content(xlsx_path, output_dir):
    base_name = os.path.basename(xlsx_path)
    txt_filename = os.path.splitext(base_name)[0] + ".txt"
    txt_path = os.path.join(output_dir, txt_filename)

    wb = load_workbook(xlsx_path, read_only=True)
    with open(txt_path, 'w', encoding='utf-8') as txt_file:
        for sheet_name in wb.sheetnames:
            txt_file.write(f"\n=== Sheet: {sheet_name} ===\n")
            sheet = wb[sheet_name]
            
            for row in sheet.iter_rows(values_only=True):
                row_text = "\t".join(str(cell) if cell is not None else "" for cell in row)
                txt_file.write(row_text + "\n")
    
    return txt_path

def process_uploads_folder():
    uploads_dir = "uploads"
    output_dir = "Youssef_Output"
    
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    if not os.path.exists(uploads_dir):
        print(f"Error: '{uploads_dir}' directory does not exist")
        return
    
    processed_files = []
    
    for filename in os.listdir(uploads_dir):
        file_path = os.path.join(uploads_dir, filename)
        
        if filename.lower().endswith('.pdf'):
            try:
                output_file = extract_pdf_content(file_path, output_dir)
                processed_files.append((filename, output_file))
            except Exception as e:
                print(f"Error processing PDF {filename}: {str(e)}")
        
        elif filename.lower().endswith('.docx'):
            try:
                output_file = extract_docx_content(file_path, output_dir)
                processed_files.append((filename, output_file))
            except Exception as e:
                print(f"Error processing DOCX {filename}: {str(e)}")
        
        elif filename.lower().endswith('.xlsx'):
            try:
                output_file = extract_xlsx_content(file_path, output_dir)
                processed_files.append((filename, output_file))
            except Exception as e:
                print(f"Error processing XLSX {filename}: {str(e)}")
    
    print("\nProcessing complete. Summary:")
    for original, output in processed_files:
        print(f"- {original} → {output}")

if __name__ == "__main__":
    process_uploads_folder()